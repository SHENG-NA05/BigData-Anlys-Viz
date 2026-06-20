from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from app.crud.curation import get_curation_theme
from app.crud.proposal import create_proposal, get_proposal as get_proposal_record
from app.crud.proposal import update_proposal as update_proposal_record
from app.db.models import Proposal
from app.db.session import get_db
from app.schemas.proposal import ProposalExportRequest, ProposalUpdateRequest
from app.services.ai_service import AIService, AIServiceError
from app.services.catalog_match_service import match_catalog_books

router = APIRouter()

# 共用 AI 服務實例（延遲初始化 Gemini 模型）
_ai_service = AIService()


@router.get("/proposals")
def list_proposals(db: Session = Depends(get_db)):
    from app.crud.proposal import list_proposals as list_proposals_record
    proposals = list_proposals_record(db)
    return {
        "status": "success",
        "data": [_proposal_to_response(p) for p in proposals],
    }


@router.get("/proposals/{proposal_id}")
def get_proposal(proposal_id: str, db: Session = Depends(get_db)):
    proposal = get_proposal_record(db, proposal_id)
    if proposal is None:
        raise HTTPException(status_code=404, detail=f"Proposal not found: {proposal_id}")

    return {
        "status": "success",
        "data": _proposal_to_response(proposal),
    }


@router.put("/proposals/{proposal_id}")
def update_proposal(
    proposal_id: str,
    request: ProposalUpdateRequest,
    db: Session = Depends(get_db),
):
    proposal = get_proposal_record(db, proposal_id)
    if proposal is None:
        raise HTTPException(status_code=404, detail=f"Proposal not found: {proposal_id}")

    try:
        update_proposal_record(db, proposal, request.title, request.content, request.status)
    except SQLAlchemyError as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to update proposal") from exc

    return {
        "status": "success",
        "data": _proposal_to_response(proposal),
    }


@router.post("/export_to_proposal")
def export_to_proposal(request: ProposalExportRequest, db: Session = Depends(get_db)):
    """
    拋轉至企劃管理中心（建立企劃草案）

    1. 驗證來源策展主題存在
    2. 呼叫 AI 擴寫企劃書草案（HTML 格式）
    3. 匹配相關館藏書籍
    4. 建立企劃書記錄存入資料庫
    """
    theme = get_curation_theme(db, request.theme_id)
    if theme is None:
        raise HTTPException(status_code=404, detail=f"Curation theme not found: {request.theme_id}")

    # 呼叫 AI 擴寫企劃書內容（HTML 格式）
    try:
        ai_content = _ai_service.expand_proposal(
            title=request.title,
            outline=request.outline,
            target_audience=request.target_audience,
        )
    except AIServiceError as exc:
        # AI 擴寫失敗時，使用基礎模板作為降級方案（Fallback）
        ai_content = _build_fallback_content(request)
        import logging
        logging.getLogger(__name__).warning(
            "AI 企劃書擴寫失敗，使用降級模板：%s", str(exc)
        )

    # 匹配館藏書籍
    matched_books = match_catalog_books(db, _build_match_keywords(theme, request))

    proposal = Proposal(
        proposal_id=_generate_proposal_id(),
        theme_id=request.theme_id,
        title=request.title,
        content=ai_content,
        matched_books=matched_books,
        status="draft",
    )

    try:
        create_proposal(db, proposal)
    except SQLAlchemyError as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to create proposal") from exc

    return {
        "status": "success",
        "proposal_id": proposal.proposal_id,
        "message": "已成功建立企劃書草案，請至企劃管理中心編輯。",
    }


def _generate_proposal_id() -> str:
    return "P" + datetime.now().strftime("%Y%m%d%H%M%S%f")


def _build_fallback_content(request: ProposalExportRequest) -> str:
    """AI 擴寫失敗時的降級方案：使用基礎 HTML 模板。"""
    return (
        f"<h1>策展宗旨與目標</h1>"
        f"<p>本次策展主題為「{request.title}」，"
        f"旨在為{request.target_audience}提供豐富的閱讀體驗。</p>"
        f"<h1>展區規劃與空間佈置</h1>"
        f"<p>{request.outline}</p>"
        f"<p>（本草案因 AI 服務暫時無法使用，僅提供基礎架構，請於編輯器中自行補充完善。）</p>"
        f"<h1>宣傳與推廣時程</h1>"
        f"<p>請依據實際展期規劃前、中、後期之宣傳活動。</p>"
        f"<h1>預算與資源評估</h1>"
        f"<p>請依據展覽規模編列相關預算。</p>"
    )


def _build_match_keywords(theme, request: ProposalExportRequest) -> list[str]:
    return [
        *(theme.keywords or []),
        theme.title,
        request.title,
        request.outline,
        request.target_audience,
    ]


def _proposal_to_response(proposal: Proposal) -> dict:
    return {
        "proposal_id": proposal.proposal_id,
        "theme_id": proposal.theme_id,
        "title": proposal.title,
        "content": proposal.content,
        "matched_books": proposal.matched_books,
        "status": proposal.status,
        "created_at": proposal.created_at.strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": proposal.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
    }


# ---------------------------------------------------------------------------
# C-4: Word & PDF 檔案產生與匯出功能
# ---------------------------------------------------------------------------
import io
from fastapi import Query
from fastapi.responses import StreamingResponse
from urllib.parse import quote
from docx import Document
from html.parser import HTMLParser


class DocxHTMLParser(HTMLParser):
    def __init__(self, document):
        super().__init__()
        self.doc = document
        self.current_tag = None
        self.list_level = 0
        self.para = None
        self.table_data = []

    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag == 'h1':
            self.para = None
        elif tag == 'h2':
            self.para = None
        elif tag == 'p':
            self.para = self.doc.add_paragraph()
        elif tag == 'ul' or tag == 'ol':
            self.list_level += 1
        elif tag == 'li':
            self.para = self.doc.add_paragraph(style='List Bullet')
        elif tag == 'table':
            self.table_data = []
        elif tag == 'tr':
            self.table_data.append([])

    def handle_endtag(self, tag):
        if tag in ('ul', 'ol'):
            self.list_level = max(0, self.list_level - 1)
        elif tag == 'table':
            if self.table_data:
                rows_count = len(self.table_data)
                cols_count = max(len(r) for r in self.table_data) if rows_count > 0 else 0
                if cols_count > 0:
                    tbl = self.doc.add_table(rows=rows_count, cols=cols_count)
                    tbl.style = 'Table Grid'
                    for r_idx, row in enumerate(self.table_data):
                        for c_idx, val in enumerate(row):
                            if c_idx < cols_count:
                                tbl.cell(r_idx, c_idx).text = val
            self.table_data = []
        self.current_tag = None

    def handle_data(self, data):
        text = data.strip()
        if not text:
            return
        if self.current_tag == 'h1':
            self.doc.add_heading(text, level=1)
        elif self.current_tag == 'h2':
            self.doc.add_heading(text, level=2)
        elif self.current_tag in ('p', 'li'):
            if self.para:
                self.para.add_run(text)
        elif self.current_tag in ('td', 'th'):
            if self.table_data:
                self.table_data[-1].append(text)
        elif not self.current_tag:
            self.doc.add_paragraph(text)


def generate_docx(proposal: Proposal) -> bytes:
    doc = Document()
    doc.add_heading(proposal.title, level=0)

    parser = DocxHTMLParser(doc)
    parser.feed(proposal.content)

    if proposal.matched_books:
        doc.add_heading("推薦匹配館藏圖書", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = '書名'
        hdr_cells[1].text = '作者'
        hdr_cells[2].text = 'ISBN'
        hdr_cells[3].text = '分類號'
        for book in proposal.matched_books:
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(book.get("title", ""))
            row_cells[1].text = str(book.get("author", ""))
            row_cells[2].text = str(book.get("isbn", ""))
            row_cells[3].text = str(book.get("classification_no", ""))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()


# PDF Generation Helpers
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import os


def register_chinese_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simsun.ttf",
    ]
    for path in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("ChineseFont", path))
                return "ChineseFont"
            except Exception:
                continue
    return "Helvetica"


class PDFHTMLParser(HTMLParser):
    def __init__(self, font_name, h1_style, h2_style, body_style):
        super().__init__()
        self.font_name = font_name
        self.h1_style = h1_style
        self.h2_style = h2_style
        self.body_style = body_style
        self.story = []
        self.current_tag = None
        self.para_text = ""
        self.table_data = []

    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag in ('h1', 'h2', 'p', 'li'):
            self.para_text = ""
        elif tag == 'table':
            self.table_data = []
        elif tag == 'tr':
            self.table_data.append([])

    def handle_endtag(self, tag):
        if tag == 'h1':
            self.story.append(Paragraph(self.para_text, self.h1_style))
            self.story.append(Spacer(1, 5))
        elif tag == 'h2':
            self.story.append(Paragraph(self.para_text, self.h2_style))
            self.story.append(Spacer(1, 4))
        elif tag in ('p', 'li'):
            prefix = "• " if tag == 'li' else ""
            self.story.append(Paragraph(prefix + self.para_text, self.body_style))
        elif tag == 'table':
            if self.table_data:
                formatted_data = []
                for row in self.table_data:
                    formatted_row = []
                    for cell in row:
                        formatted_row.append(Paragraph(cell, self.body_style))
                    formatted_data.append(formatted_row)

                t = Table(formatted_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.whitesmoke]),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ]))
                self.story.append(t)
                self.story.append(Spacer(1, 10))
        self.current_tag = None

    def handle_data(self, data):
        text = data.strip()
        if not text:
            return
        if self.current_tag in ('h1', 'h2', 'p', 'li'):
            self.para_text += text
        elif self.current_tag in ('td', 'th'):
            if self.table_data:
                self.table_data[-1].append(text)
        elif not self.current_tag:
            self.story.append(Paragraph(text, self.body_style))


def generate_pdf(proposal: Proposal) -> bytes:
    font_name = register_chinese_font()
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=A4)
    story = []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=24,
        leading=28,
        spaceAfter=15,
        alignment=1
    )
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=18,
        leading=22,
        spaceBefore=15,
        spaceAfter=10
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceBefore=10,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14,
        spaceAfter=8
    )

    story.append(Paragraph(proposal.title, title_style))
    story.append(Spacer(1, 10))

    parser = PDFHTMLParser(font_name, h1_style, h2_style, body_style)
    parser.feed(proposal.content)
    story.extend(parser.story)

    if proposal.matched_books:
        story.append(Spacer(1, 15))
        story.append(Paragraph("推薦匹配館藏圖書", h1_style))
        story.append(Spacer(1, 5))

        table_rows = [[
            Paragraph("<b>書名</b>", body_style),
            Paragraph("<b>作者</b>", body_style),
            Paragraph("<b>ISBN</b>", body_style),
            Paragraph("<b>分類號</b>", body_style)
        ]]
        for book in proposal.matched_books:
            table_rows.append([
                Paragraph(str(book.get("title", "")), body_style),
                Paragraph(str(book.get("author", "")), body_style),
                Paragraph(str(book.get("isbn", "")), body_style),
                Paragraph(str(book.get("classification_no", "")), body_style)
            ])

        t = Table(table_rows)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ]))
        story.append(t)

    doc.build(story)
    return file_stream.getvalue()


@router.get("/proposals/{proposal_id}/export")
def export_proposal(
    proposal_id: str,
    format: str = Query("docx", pattern="^(docx|pdf)$"),
    db: Session = Depends(get_db)
):
    proposal = get_proposal_record(db, proposal_id)
    if proposal is None:
        raise HTTPException(status_code=404, detail=f"Proposal not found: {proposal_id}")

    # 紀錄效益日誌 (每次匯出：省下 12.0 小時)
    try:
        from app.db.models import CostBenefitLog, SystemSetting
        rate_setting = db.query(SystemSetting).filter(SystemSetting.setting_key == "hourly_rate").first()
        hourly_rate = float(rate_setting.setting_value) if rate_setting else 200.0

        log = CostBenefitLog(
            action="proposal_export",
            target_id=proposal_id,
            time_saved_hours=12.0,
            cost_saved_amount=12.0 * hourly_rate,
        )
        db.add(log)
        db.commit()
    except Exception as exc:
        import logging
        logging.getLogger(__name__).warning("記錄匯出效益日誌失敗：%s", str(exc))

    if format == "docx":
        content = generate_docx(proposal)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{proposal.title}.docx"
    else:
        content = generate_pdf(proposal)
        media_type = "application/pdf"
        filename = f"{proposal.title}.pdf"

    encoded_filename = quote(filename)
    headers = {
        "Access-Control-Expose-Headers": "Content-Disposition",
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
    }

    return StreamingResponse(io.BytesIO(content), media_type=media_type, headers=headers)


@router.post("/proposals/{proposal_id}/match")
def match_proposal_books(proposal_id: str, db: Session = Depends(get_db)):
    proposal = get_proposal_record(db, proposal_id)
    if proposal is None:
        raise HTTPException(status_code=404, detail=f"Proposal not found: {proposal_id}")

    keywords = [proposal.title, proposal.content]
    theme = None
    if proposal.theme_id:
        theme = get_curation_theme(db, proposal.theme_id)
        if theme:
            keywords.extend(theme.keywords or [])
            keywords.append(theme.title)

    matched_books = match_catalog_books(db, keywords)
    proposal.matched_books = matched_books
    try:
        db.commit()
    except SQLAlchemyError as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to update proposal matched books") from exc

    return {
        "status": "success",
        "data": _proposal_to_response(proposal),
    }

